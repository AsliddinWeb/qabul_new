from docx import Document
from docx.shared import Inches
from PIL import Image

def create_word_template(full_name, id):
    document = Document("./static/Kvitansiya.docx")

    # Ma'lumotlarni templatega kiritish
    table1, table2 = document.tables[0], document.tables[1]
    table1.rows[3].cells[3].text = f"{full_name}"
    table2.rows[3].cells[3].text = f"{full_name}"

    # Faylni saqlash
    temp_docx = f'./media/invoyslar/invoys_{id}.docx'
    document.save(temp_docx)
    print(f"Success {id} invoys checked")

def create_malumotnoma():
    pass

# def create_malumotnoma(request):
#     doc = Document("./static/malumotnoma.docx")
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             if run.text == "{{rasm}}":
#                 run.text = ""
#                 run.add_picture(f".{request.user.image.url}", width=Inches(1.2), height=Inches(1))
#
#     doc.save(f"./media/malumotnomalar/malumotnoma_6.docx")


# def create_malumotnoma(request):
#     doc = Document("./static/malumotnoma.docx")
#     image_path = f".{request.user.image.url}"
#
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             if run.text == "{{rasm}}":
#                 run.text = ""
#                 run.add_picture(image_path, width=Inches(1), height=Inches(1))
#
#     doc.save(f"./media/malumotnomalar/malumotnoma_6.docx")
